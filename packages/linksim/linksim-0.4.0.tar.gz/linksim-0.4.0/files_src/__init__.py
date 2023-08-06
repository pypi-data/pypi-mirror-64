import xlrd
import datetime
import docx
import re
import os
import tempfile

from docx.shared import Inches
from copy import deepcopy
from contextlib import contextmanager

try:
    from PIL import Image
except Exception as e:
    print("Import PIL error!, try pip install pillow")

KEY_RE = re.compile(r'(\$\{[\w\:]+?\})')

class TaskExcel:

    def __init__(self, file_path:str, key_row=0, col_limit=30):
        self._f = file_path
        data = xlrd.open_workbook(file_path)
        table = data.sheets()[0]
        self.cases = []
        self.keys = []
        for i in range(table.nrows):
            if i == key_row:
                self.keys = [str(i.value) for no,i in enumerate(table.row(i)) if no < col_limit]
            if len(self.keys) == 0:
                continue
            d = {}
            for co in range(len(self.keys)):
                vv = table.cell(i,co)
                if vv.ctype == 3:
                    val = datetime.datetime(*xlrd.xldate_as_tuple(vv.value, 0)).ctime()
                else:
                    val = vv.value
                d[self.keys[co]] = val

            self.cases.append(d)
            
    def render(self, redner_func, default="case.html"):
        return redner_func(default, cases=self.cases)

class Word:
    def __init__(self, file_path):
        self._doc = docx.Document(file_path)
        self._keys = []
        self._reverse = {}
        self._map = {}
        self._old_doc = deepcopy(self._doc)
        self._data = {}
        self.scan()

    def keys(self):
        return self._keys
    
    def scan(self):
        paragraphs = []
        self._keys = []
        self._map = {}
        for t in self._doc.tables:
            for row in t.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        paragraphs.append(paragraph)
                        if '${' in paragraph.text:
                            for k in KEY_RE.findall(paragraph.text):
                                self._keys.append(k)


        for paragraph in self._doc.paragraphs:
            paragraphs.append(paragraph)
            if '${' in paragraph.text:
                for k in KEY_RE.findall(paragraph.text):
                    self._keys.append(k)
        
        for k in self._keys:
            d = {}
            if ':' in k:
                label, name = k.split(":",1)
                label = label[2:]
                name = name[:-1]
            else:
                label = "textarea"
                name = k[2:-1]
            d["name"] = name
            d["type"] = label
            self._map[k] = d
            self._reverse[name] = k
        
        return paragraphs

    def docx_render(self, saved_name, data):
        if len(data) == 0:return
        # if set(data.keys()) != set(self._reverse.values()):
        #     print(">>>\n",data.keys(), "\n", self._reverse.values())
        #     return
        data = data
        paragraphs = self.scan()

        for p in paragraphs:
            for key, val in data.items():
                key_name = '${{{}}}'.format(key) # I'm using placeholders in the form ${PlaceholderName}
                
                if key_name in p.text:
                    inline = p.runs
                    # Replace strings and retain the same style.
                    # The text to be replaced can be split over several runs so
                    # search through, identify which runs need to have text replaced
                    # then replace the text in those identified
                    started = False
                    key_index = 0
                    # found_runs is a list of (inline index, index of match, length of match)
                    found_runs = list()
                    found_all = False
                    replace_done = False
                    # insert_image = False
                    if key_name.startswith("${img:"):
                        img_run = p.add_run()
                        if isinstance(val, list):
                            for v in val:    
                                if os.path.exists(v) and v.split(".")[-1] in ('png','jpg'):
                                    try:
                                        img_run.add_picture(v,width=Inches(4.0))
                                    except docx.image.exceptions.UnrecognizedImageError:
                                        print("add picture failed:", v, "try convert to png")
                                        img = Image.open(v)
                                        vpng = v.split(".")[:-1] + ".png"
                                        img.save(vpng)
                                        print("try add_picture by png...")
                                        img_run.add_picture(vpng)
                        # insert_image = True
                        continue

                    for i in range(len(inline)):

                        # case 1: found in single run so short circuit the replace
                        if key_name in inline[i].text and not started:
                            found_runs.append((i, inline[i].text.find(key_name), len(key_name)))
                            text = inline[i].text.replace(key_name, str(val))
                            inline[i].text = text
                            replace_done = True
                            found_all = True
                            break

                        if key_name[key_index] not in inline[i].text and not started:
                            # keep looking ...
                            continue

                        # case 2: search for partial text, find first run
                        if key_name[key_index] in inline[i].text and inline[i].text[-1] in key_name and not started:
                            # check sequence
                            start_index = inline[i].text.find(key_name[key_index])
                            check_length = len(inline[i].text)
                            for text_index in range(start_index, check_length):
                                if inline[i].text[text_index] != key_name[key_index]:
                                    # no match so must be false positive
                                    break
                            if key_index == 0:
                                started = True
                            chars_found = check_length - start_index
                            key_index += chars_found
                            found_runs.append((i, start_index, chars_found))
                            if key_index != len(key_name):
                                continue
                            else:
                                # found all chars in key_name
                                found_all = True
                                break

                        # case 2: search for partial text, find subsequent run
                        if key_name[key_index] in inline[i].text and started and not found_all:
                            # check sequence
                            chars_found = 0
                            check_length = len(inline[i].text)
                            for text_index in range(0, check_length):
                                if inline[i].text[text_index] == key_name[key_index]:
                                    key_index += 1
                                    chars_found += 1
                                else:
                                    break
                            # no match so must be end
                            found_runs.append((i, 0, chars_found))
                            if key_index == len(key_name):
                                found_all = True
                                break

                    if found_all and not replace_done:
                        for i, item in enumerate(found_runs):
                            index, start, length = [t for t in item]
                            if i == 0:
                                text = inline[index].text.replace(inline[index].text[start:start + length], str(val))
                                inline[index].text = text
                            else:
                                text = inline[index].text.replace(inline[index].text[start:start + length], '')
                                inline[index].text = text
                    # print(p.text)
        self._doc.save(saved_name)
        self._doc = deepcopy(self._old_doc)

    def render_form(self,redner_func, template_name="defualt_form.html", action="/word"):
        
        for k in self._keys:
            d = {}
            if ':' in k:
                label, name = k.split(":",1)
                label = label[2:]
                name = name[:-1]
            else:
                label = "textarea"
                name = k[2:-1]
            d["name"] = name
            d["type"] = label
            self._map[k] = d
            self._reverse[name] = k
        
        return redner_func(template_name, forms=self._map, action=action)
    
    @contextmanager
    def parse_from_flask_request(self, request):
        if len(self._reverse) == 0:
            self.scan()
        try:
            form = request.form
            files = request.files
            with tempfile.TemporaryDirectory() as tmpdir:
                data = {}
                reverse_map = self._reverse
                for k in files.keys():
                    if k not in reverse_map:
                        print("not found in map:",k)
                        continue
                    f_handlers = files.getlist(k)
                    data[reverse_map[k][2:-1]] = []
                    for f_handler in f_handlers:
                        val = os.path.join(tmpdir, f_handler.filename)
                        f_handler.save(val)
                        data[reverse_map[k][2:-1]].append(val)
                        print("cache->", val)

                for k in form.keys():
                    if k not in reverse_map:
                        print("not found in map:",k)
                        continue
                    
                    val = form[k]
                    data[reverse_map[k][2:-1]] = val
                yield data
        except Exception as e:
            raise e
        finally:
            pass